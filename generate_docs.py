"""生成用户手册和工具介绍 Word 文档"""

import docx
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os


def set_cell_shading(cell, color):
    """给表格单元格设置背景色"""
    shading_elm = docx.oxml.OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_paragraph(doc, text, style_name=None, bold=False, size=None, color=None, alignment=None, space_after=None, space_before=None):
    """添加带样式的段落"""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_heading_styled(doc, text, level=1):
    """添加标题"""
    h = doc.add_heading(text, level=level)
    return h


def add_bullet(doc, text, bold_prefix="", level=0):
    """添加项目符号段落"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_note_box(doc, text):
    """添加提示段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("💡 提示：")
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(100, 100, 100)
    return p


def add_code_block(doc, text):
    """添加等宽字体的代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p


# ═══════════════════════════════════════════
# 文档1: 用户手册
# ═══════════════════════════════════════════

def generate_user_manual():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── 封面标题 ──
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("讯飞语音转写工具 v1.0")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 90, 160)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("用 户 手 册")
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run("适用于 XfyunTranscriber.exe   |   2025年6月")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    # ── 目录页 ──
    add_heading_styled(doc, "目录", level=1)
    toc_items = [
        "一、启动程序与界面概览",
        "二、配置 API 凭证",
        "三、选择音频文件与信息查看",
        "四、设置转写参数",
        "    4.1 说话人分离",
        "    4.2 音频裁剪",
        "    4.3 自动转换为 WAV",
        "五、启动转写与进度监控",
        "六、查看与保存结果",
        "七、测试凭证",
        "八、常见问题",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item).font.size = Pt(11)
    doc.add_page_break()

    # ════════════════════════════════
    # 一、启动程序与界面概览
    # ════════════════════════════════
    add_heading_styled(doc, "一、启动程序与界面概览", level=1)

    add_styled_paragraph(doc, "1.1 如何打开程序", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("双击 ").font.size = Pt(10.5)
    run = p.add_run("XfyunTranscriber.exe")
    run.bold = True
    run.font.size = Pt(10.5)
    p.add_run(" 文件即可启动程序。首次启动需要 5-15 秒的解压时间，请耐心等待。").font.size = Pt(10.5)

    add_note_box(doc, "该程序无需安装 Python 或其他运行环境，下载后即可直接使用。运行时需要保持网络连接，以便访问讯飞云端转写服务。")

    add_styled_paragraph(doc, "1.2 主界面布局", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("程序主窗口从上到下依次包含以下七个区域：").font.size = Pt(10.5)

    # 区域说明表格
    table = doc.add_table(rows=8, cols=3, style='Table Grid')
    table.autofit = True

    headers = ["区域编号", "区域名称", "功能说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(10)

    rows_data = [
        ["①", "API 设置", "填写讯飞服务凭证，配置说话人数量，保存/验证凭证"],
        ["②", "音频文件", "选择本地音频文件，查看文件时长、格式等信息"],
        ["③", "裁剪设置（可选）", "设置音频裁剪的起始时间和时长"],
        ["④", "转写选项", "选择是否转换为 WAV 格式、是否开启说话人分离"],
        ["⑤", "操作按钮", "开始转写 / 取消当前任务"],
        ["⑥", "进度区域", "显示转写进度条、当前状态和等待时间"],
        ["⑦", "结果区域", "显示转写文本结果，可保存或复制"],
    ]
    for i, row_data in enumerate(rows_data):
        for j, text in enumerate(row_data):
            table.rows[i+1].cells[j].text = text
            for p in table.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

    # ════════════════════════════════
    # 二、配置 API 凭证
    # ════════════════════════════════
    add_heading_styled(doc, "二、配置 API 凭证", level=1)

    add_styled_paragraph(doc, "2.1 凭证说明", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("程序需要两个必填凭证才能调用讯飞语音转写服务：").font.size = Pt(10.5)

    add_bullet(doc, "（APPID）：讯飞开放平台应用的唯一标识，8位纯数字", bold_prefix="APPID")
    add_bullet(doc, "（APISecret）：用于 API 请求签名的密钥，字母和数字组成", bold_prefix="APISecret")
    add_bullet(doc, "（APIKey）：可选字段，当前版本不参与 API 调用，留空即可", bold_prefix="APIKey")

    add_styled_paragraph(doc, "2.2 如何获取凭证", bold=True, size=11)
    steps = [
        "打开浏览器，访问讯飞开放平台：https://www.xfyun.cn/",
        "注册并登录账号（如已有账号可直接登录）",
        "进入控制台，点击「创建新应用」，填写应用名称和描述",
        "在应用列表中找到刚创建的应用，点击进入详情",
        "在「语音识别」分类下，找到并开通「录音文件识别」服务",
        "在应用详情页的「应用信息」中，即可看到 APPID 和 APISecret",
    ]
    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"第{i+1}步：")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(step).font.size = Pt(10)

    add_note_box(doc, "请务必开通「录音文件识别」服务，否则即使凭证正确也无法正常转写。")

    add_styled_paragraph(doc, "2.3 填写和保存凭证", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("在程序顶部「API 设置」区域，按以下步骤操作：").font.size = Pt(10.5)

    sub_steps = [
        "在「APPID」输入框中填入从讯飞平台获取的 APPID（8位数字）",
        "「APIKey」输入框可留空（当前版本不使用此字段）",
        "在「APISecret」输入框中填入从讯飞平台获取的 APISecret（输入时显示为 ****）",
        "点击右侧「验证凭证」按钮，确认凭证格式是否正确",
        "点击「保存设置」按钮，凭证将被保存到本地，下次启动自动加载",
    ]
    for i, step in enumerate(sub_steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"步骤 {i+1}：")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(step).font.size = Pt(10)

    add_note_box(doc, "凭证保存在本地计算机中（%APPDATA%\\xfyun_transcriber\\config.json），不会上传到任何第三方服务器。")

    # ════════════════════════════════
    # 三、选择音频文件与信息查看
    # ════════════════════════════════
    add_heading_styled(doc, "三、选择音频文件与信息查看", level=1)

    add_styled_paragraph(doc, "3.1 支持的音频格式", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("程序支持以下常见音频格式：").font.size = Pt(10.5)

    formats = ["MP3（.mp3）", "WAV（.wav）", "FLAC（.flac）", "M4A（.m4a）",
               "AAC（.aac）", "OGG（.ogg）", "WMA（.wma）", "AMR（.amr）"]
    for fmt in formats:
        add_bullet(doc, fmt)

    add_note_box(doc, "最大支持 500MB 的音频文件。如需处理更大文件，请先使用音频编辑软件分割后再导入。")

    add_styled_paragraph(doc, "3.2 选择文件", bold=True, size=11)
    steps = [
        "在「音频文件」区域，点击右侧的「选择文件」按钮",
        "在弹出的文件浏览窗口中，找到并选中您的音频文件",
        "点击「打开」确认选择",
    ]
    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"操作 {i+1}：")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(step).font.size = Pt(10)

    add_styled_paragraph(doc, "3.3 查看文件信息", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("选择文件后，文件路径下方会显示该音频的基本信息：").font.size = Pt(10.5)

    add_bullet(doc, "时长：音频的总播放时长（格式：时:分:秒）")
    add_bullet(doc, "格式：音频文件的格式（如 MP3、WAV 等）")
    add_bullet(doc, "采样率：音频的采样率（单位：Hz）")
    add_bullet(doc, "声道：单声道或立体声")
    add_bullet(doc, "大小：文件占用空间（单位：MB）")

    add_note_box(doc, "如果提示「音频详情需要 ffmpeg」，说明程序内置的 ffmpeg 组件未正确加载，请联系技术支持。")

    # ════════════════════════════════
    # 四、设置转写参数
    # ════════════════════════════════
    add_heading_styled(doc, "四、设置转写参数", level=1)

    # 4.1 说话人分离
    add_styled_paragraph(doc, "4.1 说话人分离（声纹区分）", bold=True, size=13)
    p = doc.add_paragraph()
    p.add_run("在「API 设置」区域底部和「转写选项」区域，有两个与说话人分离相关的设置：").font.size = Pt(10.5)

    add_bullet(doc, "说话人数量：点击下拉框，可选择「自动」或指定 2-10 人", bold_prefix="① ")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run("自动")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run("：由讯飞 AI 自动判断音频中有几位说话人").font.size = Pt(10)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run("手动（2-10）")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run("：明确告诉 AI 音频中有几位说话人，适合已知人数的场景").font.size = Pt(10)

    add_bullet(doc, "在「转写选项」区域，勾选「启用说话人分离（声纹区分）」以开启该功能", bold_prefix="② ")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.add_run("建议保持勾选，这样转写结果中会标注每句话是「说话人几」说的。").font.size = Pt(10)

    # 4.2 音频裁剪
    add_styled_paragraph(doc, "4.2 音频裁剪（可选）", bold=True, size=13)
    p = doc.add_paragraph()
    p.add_run("如果只需转写音频中的某一段，可使用裁剪功能：").font.size = Pt(10.5)

    sub_steps = [
        "在「裁剪设置」区域，勾选「启用裁剪」复选框",
        "在「开始时间(秒)」输入框中，填入裁剪的起始位置",
        "在「裁剪时长(秒)」输入框中，填入需要转写的时长（留空则表示裁剪到文件结尾）",
    ]
    for i, step in enumerate(sub_steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"操作 {i+1}：")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(step).font.size = Pt(10)

    add_styled_paragraph(doc, "时间格式说明：", bold=True, size=10)
    add_bullet(doc, "纯秒数，例如 90 表示从第 90 秒开始")
    add_bullet(doc, "分:秒，例如 1:30 表示从 1 分 30 秒开始")
    add_bullet(doc, "时:分:秒，例如 0:01:30 表示从 1 分 30 秒开始")

    # 4.3 转换为WAV
    add_styled_paragraph(doc, "4.3 自动转换为 WAV", bold=True, size=13)
    p = doc.add_paragraph()
    p.add_run("在「转写选项」区域，勾选「转换为 WAV（推荐，提升说话人识别质量）」。").font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.add_run("开启后，程序会自动将音频转换为 16kHz 采样率、16bit 位深、单声道的 WAV 格式后再上传。这种格式是讯飞语音识别的最佳输入格式，能够显著提升转写准确率和说话人分离效果。").font.size = Pt(10.5)

    add_note_box(doc, "强烈建议保持此项勾选。仅在您确认原始音频已符合 16kHz/16bit/单声道 WAV 格式时，才可取消勾选以节省处理时间。")

    # ════════════════════════════════
    # 五、启动转写与进度监控
    # ════════════════════════════════
    add_heading_styled(doc, "五、启动转写与进度监控", level=1)

    add_styled_paragraph(doc, "5.1 开始转写", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("确认凭证已填写、音频文件已选择、参数已设置完毕后，点击中间的「开始转写」按钮。").font.size = Pt(10.5)

    add_note_box(doc, "点击前请确保计算机已连接互联网。转写全程需要与讯飞云端服务器通信。")

    add_styled_paragraph(doc, "5.2 转写流程", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("点击「开始转写」后，程序会自动执行以下步骤：").font.size = Pt(10.5)

    flow_steps = [
        "正在处理音频 — 如果开启了裁剪或 WAV 转换，程序会先对音频进行预处理",
        "正在上传音频 — 将处理后的音频文件上传到讯飞云端服务器",
        "转写进行中 — 讯飞 AI 对音频进行语音识别和说话人分离",
        "正在获取结果 — 从服务器下载转写完成的文本",
        "转写完成 — 结果展示在下方文本框中",
    ]
    for i, step in enumerate(flow_steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"第{i+1}步 ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(step).font.size = Pt(10)

    add_styled_paragraph(doc, "5.3 进度监控", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("在「进度」区域，您可以实时查看：").font.size = Pt(10.5)

    add_bullet(doc, "进度条：以百分比形式直观展示整体进度")
    add_bullet(doc, "状态文字：显示当前所处阶段（如「准备中…」「音频处理中…」「语音转写中…」）")
    add_bullet(doc, "已等待时间：从点击「开始转写」起累计的等待秒数")
    add_bullet(doc, "任务 ID：唯一标识本次转写任务，可用于问题排查")

    add_note_box(doc, "转写时间取决于音频时长和网络状况，一般情况下 1 小时的音频约需 5-10 分钟完成转写。任务最长等待时间为 10 分钟，超时会自动提示，任务 ID 可用于在讯飞控制台手动查看。")

    add_styled_paragraph(doc, "5.4 取消转写", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("如果需要中止当前转写，点击「取消」按钮即可。取消后可以重新选择文件再次开始。").font.size = Pt(10.5)

    # ════════════════════════════════
    # 六、查看与保存结果
    # ════════════════════════════════
    add_heading_styled(doc, "六、查看与保存结果", level=1)

    add_styled_paragraph(doc, "6.1 查看转写结果", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("转写完成后，结果会自动显示在底部「转写结果」区域的文本框中。").font.size = Pt(10.5)

    add_styled_paragraph(doc, "6.2 理解输出格式", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("转写结果采用以下格式展示：").font.size = Pt(10.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("00:00:15 - 00:00:28 说话人1")
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0, 90, 160)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("大家好，欢迎参加今天的会议。")
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("00:00:30 - 00:00:45 说话人2")
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0, 90, 160)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("谢谢大家，今天我们来讨论一下项目进度。")
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("格式说明：").font.size = Pt(10.5)
    add_bullet(doc, "第一行：时间段（开始时间 - 结束时间）+ 说话人标签")
    add_bullet(doc, "第二行：该说话人在该时间段内说的内容")
    add_bullet(doc, "空行：分隔不同的说话人段落")
    add_bullet(doc, "说话人按发言顺序编号（说话人1、说话人2 …），程序会自动合并同一说话人的连续片段")

    add_styled_paragraph(doc, "6.3 保存结果", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("程序提供三种方式处理结果：").font.size = Pt(10.5)

    add_bullet(doc, "自动保存：转写完成后，程序会自动在与原始音频相同的文件夹中生成一个 .txt 文件（文件名为「原文件名_transcribed.txt」）")
    add_bullet(doc, "手动保存：点击「保存结果」按钮，选择保存位置和文件名")
    add_bullet(doc, "复制到剪贴板：点击「复制到剪贴板」按钮，将全部结果复制到系统剪贴板，可粘贴到 Word、记事本等任意程序中")

    add_note_box(doc, "所有结果文件均使用 UTF-8 编码保存，兼容 Windows 记事本、Word、WPS 等主流文本编辑软件。")

    # ════════════════════════════════
    # 七、测试凭证
    # ════════════════════════════════
    add_heading_styled(doc, "七、测试凭证", level=1)

    add_note_box(doc, "以下凭证仅供初步功能测试使用，可能存在用量限制或有效期。正式使用请务必注册自己的讯飞账号并获取专属凭证。")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run("测试用 APPID：").bold = True
    p.add_run("f0e971cfd4b40e843f3f63a93eb7f213")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.add_run("测试用 APISecret：").bold = True
    p.add_run("OTVkN2IzZWRmZjUyY2U1YjhhNzBmZDI3")

    p = doc.add_paragraph()
    p.add_run("APIKey：").bold = True
    p.add_run("留空即可（当前版本不使用此字段）")

    doc.add_paragraph()
    add_styled_paragraph(doc, "获取正式凭证的步骤：", bold=True, size=11)
    steps = [
        "访问 https://www.xfyun.cn/ 注册账号",
        "在控制台创建新应用",
        "开通「录音文件识别」服务",
        "复制 APPID 和 APISecret 到程序中",
    ]
    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"{i+1}. ")
        run.font.size = Pt(10)
        p.add_run(step).font.size = Pt(10)

    # ════════════════════════════════
    # 八、常见问题
    # ════════════════════════════════
    add_heading_styled(doc, "八、常见问题", level=1)

    qa = [
        ("Q：点击「开始转写」后提示「APPID不能为空」？",
         "A：请先在「API 设置」区域填写 APPID 和 APISecret，确保两项都不为空。"),
        ("Q：提示「APPID应为纯数字」？",
         "A：APPID 必须为纯数字，请检查是否复制了多余的空格或字符。去讯飞控制台 → 应用详情中重新复制。"),
        ("Q：提示「签名校验失败」或「AccessKeyId 不存在」？",
         "A：请检查 APISecret 是否填写正确，确认是「录音文件识别」服务的凭证，且服务已开通。"),
        ("Q：转写进度卡住不动了？",
         "A：这是正常现象。讯飞云端处理大文件时需要一定时间，程序每 5 秒轮询一次进度。建议耐心等待，1 小时音频约需 5-10 分钟。"),
        ("Q：转写完成后没有显示说话人标签？",
         "A：请确认「说话人分离」选项已勾选。如果音频中确实只有一个人说话，或者说话人声音特征差别不大，分离效果会受影响。"),
        ("Q：转写准确率不高？",
         "A：建议保持「转换为 WAV」选项开启；使用清晰录音环境；避免背景噪音过大的音频。"),
        ("Q：程序闪退？",
         "A：请查看 %APPDATA%\\xfyun_transcriber\\app.log 日志文件，将日志发给技术支持排查。"),
    ]
    for question, answer in qa:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(question)
        run.bold = True
        run.font.size = Pt(10.5)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.add_run(answer).font.size = Pt(10)

    # ── 保存 ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "讯飞语音转写工具_用户手册.docx")
    doc.save(output_path)
    print(f"用户手册已生成: {output_path}")
    return output_path


# ═══════════════════════════════════════════
# 文档2: 工具介绍说明
# ═══════════════════════════════════════════

def generate_introduction():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── 封面 ──
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("讯飞语音转写工具 v1.0")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 90, 160)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("产品介绍与技术说明")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.add_run("一款基于讯飞人工智能的桌面端语音转文字工具").font.size = Pt(11)

    doc.add_page_break()

    # ── 1. 产品概述 ──
    add_heading_styled(doc, "一、产品概述", level=1)

    p = doc.add_paragraph()
    p.add_run("讯飞语音转写工具").bold = True
    p.add_run("是一款 Windows 桌面应用程序，能够将录音文件（如会议录音、访谈录音、课程录音等）自动转换为文字，并智能区分不同说话人。").font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.add_run("用户只需双击打开程序，填入讯飞服务凭证，选择音频文件，点击「开始转写」即可在几分钟内获得带有说话人标签和时间戳的结构化文本。程序已打包为单文件可执行程序（exe），无需安装 Python 或其他依赖，下载即用。").font.size = Pt(10.5)

    # ── 2. 适用场景 ──
    add_heading_styled(doc, "二、适用场景", level=1)

    scenarios = [
        ("会议记录", "将会议录音转为会议纪要，自动标注每位发言人的发言内容和时间点，告别手动整理录音的繁琐工作。"),
        ("访谈/采访", "新闻采访、用户调研、深度访谈等场景，快速产出带说话人区分的访谈文字稿。"),
        ("课程/讲座", "将课堂录音、讲座音频转为文字笔记，方便复习和检索关键内容。"),
        ("电话录音", "客服通话、销售回访等电话录音的批量转写和说话人分离。"),
        ("视频字幕", "为视频内容生成带有时间戳的文字稿，作为字幕制作的基础素材。"),
        ("法律/医疗", "庭审录音、医患沟通录音等需要精确文字记录的专业场景（注意：需确保符合相关隐私合规要求）。"),
    ]

    table = doc.add_table(rows=len(scenarios) + 1, cols=2, style='Table Grid')
    table.autofit = True
    table.columns[0].width = Cm(3)
    table.columns[1].width = Cm(12)

    for i, h in enumerate(["场景", "说明"]):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(10)

    for i, (scene, desc_text) in enumerate(scenarios):
        table.rows[i+1].cells[0].text = scene
        table.rows[i+1].cells[1].text = desc_text
        for j in range(2):
            for p in table.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

    # ── 3. 核心功能 ──
    add_heading_styled(doc, "三、核心功能", level=1)

    features = [
        ("语音转文字", "基于讯飞深度学习语音识别引擎，将音频中的语音内容自动转为文字，支持中文普通话。"),
        ("说话人分离", "通过声纹识别技术自动区分不同说话人，在结果中用「说话人1」「说话人2」等标签标注，支持自动判断人数或手动指定。"),
        ("音频裁剪", "支持对长音频的任意时间段进行裁剪，只转写需要的片段，节省时间和调用额度。"),
        ("格式自动优化", "自动将各类音频格式转换为 16kHz/16bit/单声道 WAV 格式，达到讯飞 API 最佳输入规格，提升识别准确率。"),
        ("多格式支持", "支持 MP3、WAV、FLAC、M4A、AAC、OGG、WMA、AMR 等常见音频格式，最大支持 500MB 文件。"),
        ("结构化输出", "输出结果包含精确的时间戳（时:分:秒）和说话人标注，段落清晰、易于阅读和后续编辑。"),
        ("断点友好", "转写过程中可随时取消，凭证可保存后自动加载，适合反复使用的工作流程。"),
    ]

    for title_text, desc_text in features:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(f"▸ {title_text}")
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.add_run(desc_text).font.size = Pt(10)

    # ── 4. 技术栈 ──
    add_heading_styled(doc, "四、技术栈", level=1)

    add_styled_paragraph(doc, "4.1 编程语言与框架", bold=True, size=12)

    techs = [
        ("Python 3.x", "编程语言", "核心开发语言，用于实现业务逻辑、API 调用、音频处理和 GUI 界面"),
        ("Tkinter", "GUI框架", "Python 标准 GUI 库，构建跨平台的图形用户界面，无需额外安装"),
        ("pydub", "音频处理", "音频处理库，提供音频格式转换、裁剪、信息读取等功能"),
        ("requests", "网络通信", "HTTP 客户端库，用于与讯飞 REST API 进行网络通信"),
        ("PyInstaller", "打包工具", "打包工具，将 Python 项目编译为独立的 Windows .exe 可执行文件"),
    ]

    table2 = doc.add_table(rows=len(techs) + 1, cols=3, style='Table Grid')
    table2.autofit = True
    table2.columns[0].width = Cm(3.5)
    table2.columns[1].width = Cm(2.5)
    table2.columns[2].width = Cm(9)

    for i, h in enumerate(["技术", "类别", "作用"]):
        cell = table2.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(10)

    for i, (name, cat, desc_text) in enumerate(techs):
        table2.rows[i+1].cells[0].text = name
        table2.rows[i+1].cells[1].text = cat
        table2.rows[i+1].cells[2].text = desc_text
        for j in range(3):
            for p in table2.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

    add_styled_paragraph(doc, "4.2 外部依赖", bold=True, size=12)
    p = doc.add_paragraph()
    p.add_run("ffmpeg / ffprobe").bold = True
    p.add_run("：开源音视频处理工具，程序内置了 ffmpeg.exe 和 ffprobe.exe（约 120MB），打包在 exe 内部，用户无需单独安装。用于音频格式转换、裁剪、信息提取等底层操作。").font.size = Pt(10.5)

    add_styled_paragraph(doc, "4.3 讯飞语音识别 API", bold=True, size=12)
    p = doc.add_paragraph()
    p.add_run("程序调用讯飞开放平台提供的").font.size = Pt(10.5)
    p.add_run("「录音文件识别（raasr）」REST API").bold = True
    p.add_run("，该 API 专为长音频文件的离线转写设计。").font.size = Pt(10.5)

    add_styled_paragraph(doc, "API 接口说明：", bold=True, size=11)

    api_table = doc.add_table(rows=4, cols=3, style='Table Grid')
    api_table.autofit = True
    for i, h in enumerate(["接口", "HTTP 方法", "说明"]):
        cell = api_table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
                run.font.size = Pt(10)

    api_data = [
        ["/v2/api/upload", "POST", "上传音频文件，返回任务 ID"],
        ["/v2/api/getProgress", "GET", "查询转写进度"],
        ["/v2/api/getResult", "GET", "获取转写结果"],
    ]
    for i, row_data in enumerate(api_data):
        for j, text in enumerate(row_data):
            api_table.rows[i+1].cells[j].text = text
            for p in api_table.rows[i+1].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()

    add_styled_paragraph(doc, "认证机制：", bold=True, size=11)
    p = doc.add_paragraph()
    p.add_run("API 采用 ").font.size = Pt(10.5)
    p.add_run("HMAC-SHA256 签名认证").bold = True
    p.add_run("。每次请求时，客户端使用 APISecret 对请求参数（包含 APPID 和时间戳）进行签名，服务器端用相同的密钥验证签名，确保请求来源合法、数据未被篡改。").font.size = Pt(10.5)

    # ── 5. 架构设计 ──
    add_heading_styled(doc, "五、架构设计", level=1)

    p = doc.add_paragraph()
    p.add_run("程序采用模块化设计，各模块职责分明：").font.size = Pt(10.5)

    modules = [
        ("main.py", "应用入口，初始化日志系统和 Tkinter 主窗口"),
        ("gui.py", "GUI 界面层，处理用户交互、按钮事件、进度显示"),
        ("api_client.py", "API 通信层，封装讯飞 raasr API 的签名、上传、查询、获取结果"),
        ("audio_processor.py", "音频处理层，封装 ffmpeg/pydub 的格式转换和裁剪操作"),
        ("poll_manager.py", "轮询管理层，管理转写进度的定时轮询和超时/取消逻辑"),
        ("result_formatter.py", "结果格式化层，解析 API 返回的 JSON 数据，合并说话人段落，格式化输出"),
        ("config.py", "配置管理层，凭证的加载、保存和格式校验"),
        ("utils.py", "工具函数层，时间格式化、路径处理、日志配置等通用功能"),
        ("exceptions.py", "异常体系，定义各类业务异常和错误码映射"),
    ]

    for name, desc_text in modules:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"{name}  —  ")
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.bold = True
        p.add_run(desc_text).font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("处理流程如下：").font.size = Pt(10.5)
    flow = [
        "用户通过 GUI 选择文件、设置参数 →",
        "音频处理层（pydub + ffmpeg）进行裁剪/格式转换 →",
        "API 通信层上传音频文件至讯飞云端 →",
        "轮询管理层定时检查转写进度 →",
        "转写完成后获取原始 JSON 结果 →",
        "结果格式化层解析并合并说话人段落 →",
        "结构化文本展示在 GUI 界面，同时自动保存为 TXT 文件",
    ]
    for i, f in enumerate(flow):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"{i+1}. ")
        run.font.size = Pt(10)
        p.add_run(f).font.size = Pt(10)

    # ── 6. 技术选型理由 ──
    add_heading_styled(doc, "六、技术选型理由", level=1)

    selections = [
        ("为什么选择讯飞 raasr API？",
         "讯飞（科大讯飞）是国内语音识别领域的领先厂商，其录音文件识别 API 具有以下优势：\n"
         "  • 识别准确率高：基于深度学习模型，对中文普通话的识别准确率在业界领先\n"
         "  • 原生支持说话人分离：API 内置声纹分离功能，无需额外处理\n"
         "  • 支持大文件：单文件最大支持 500MB（约 30 小时音频）\n"
         "  • 调用成本合理：免费额度 + 按量付费，适合个人和小团队使用\n"
         "  • 稳定可靠：讯飞云端服务具备高可用架构，API 响应稳定"),
        ("为什么选择 Python + Tkinter？",
         "Python 拥有丰富的第三方库生态，pydub、requests 等库大大降低了音频处理和 HTTP 通信的开发成本。Tkinter 作为 Python 标准库自带的 GUI 框架，无需额外安装，打包体积小，适合开发轻量级桌面工具。"),
        ("为什么使用 PyInstaller 打包？",
         "PyInstaller 能将 Python 程序及其所有依赖（包括 ffmpeg.exe）打包成单个 exe 文件，用户无需安装 Python 环境或任何依赖库，双击即可运行，极大降低了非技术用户的使用门槛。"),
        ("为什么内置 ffmpeg？",
         "ffmpeg 是音频处理的事实标准工具。将其内置到 exe 中，用户无需单独下载和配置 ffmpeg，避免了环境配置的复杂性。程序通过 monkey-patch 机制确保 pydub 在打包模式下仍能正确定位 ffmpeg。"),
    ]

    for title_text, content in selections:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(content).font.size = Pt(10)

    # ── 7. 部署与分发 ──
    add_heading_styled(doc, "七、部署与分发", level=1)

    p = doc.add_paragraph()
    p.add_run("程序打包为单文件可执行程序，部署方式极为简单：").font.size = Pt(10.5)

    add_bullet(doc, "输出文件：XfyunTranscriber.exe（约 50.9 MB）")
    add_bullet(doc, "运行环境：Windows 7 及以上操作系统")
    add_bullet(doc, "网络要求：运行时需连接互联网（访问讯飞 API）")
    add_bullet(doc, "无需安装：双击即可运行，首次启动需 5-15 秒解压时间")
    add_bullet(doc, "凭证持久化：用户配置的 APPID/APISecret 保存在 %APPDATA%\\xfyun_transcriber\\config.json")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("构建命令：").bold = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run("pyinstaller build.spec")
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    # ── 8. 注意事项 ──
    add_heading_styled(doc, "八、注意事项", level=1)

    notes = [
        "本工具依赖讯飞开放平台的云端服务，使用前需注册讯飞账号并开通「录音文件识别」服务。",
        "API 调用可能产生费用，请关注讯飞控制台的用量和计费规则。",
        "音频文件通过 HTTPS 加密上传至讯飞服务器，传输过程安全。但对涉密或敏感内容的音频，请评估是否符合所在组织的信息安全政策。",
        "程序内置的 ffmpeg 来源于 FFmpeg 项目（https://ffmpeg.org/），遵循 LGPL/GPL 许可证。",
        "Windows 防火墙或杀毒软件可能在首次运行时拦截程序，请允许其访问网络。",
    ]

    for note in notes:
        add_bullet(doc, note)

    # ── 保存 ──
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "讯飞语音转写工具_产品介绍.docx")
    doc.save(output_path)
    print(f"产品介绍已生成: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_user_manual()
    print("---")
    generate_introduction()
    print("全部文档生成完成！")
