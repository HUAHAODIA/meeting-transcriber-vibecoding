"""生成说话人分离Bug排除报告 Word 文档"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import docx


def set_cell_shading(cell, color):
    shading_elm = docx.oxml.OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    shading_elm.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_heading_styled(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_styled_paragraph(doc, text, bold=False, size=11, color=None, alignment=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
    p.add_run(text).font.size = Pt(10.5)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p


def add_note(doc, text, label="提示"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"💡 {label}：")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(100, 100, 100)
    return p


def generate():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── 封面 ──
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("讯飞语音转写工具 v1.0")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0, 90, 160)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Bug 排除报告")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    topic = doc.add_paragraph()
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = topic.add_run("问题：说话人分离未生效 — 转写结果仅显示1位说话人")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(180, 60, 0)

    doc.add_paragraph()
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver.add_run("报告日期：2025年6月4日")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    # ════════════════════════════════
    # 一、问题概述
    # ════════════════════════════════
    add_heading_styled(doc, "一、问题概述", level=1)

    add_styled_paragraph(doc, "1.1 问题描述", bold=True, size=12)
    p = doc.add_paragraph()
    p.add_run("用户反馈：运行 XfyunTranscriber.exe 进行语音转写时，在 GUI 界面「说话人数量」下拉框中选择了「自动」，但转写完成后所有文本均被标记为「说话人1」，未能识别出实际音频中的多位说话人。").font.size = Pt(10.5)

    add_styled_paragraph(doc, "1.2 测试环境", bold=True, size=12)

    tbl = doc.add_table(rows=5, cols=2, style='Table Grid')
    tbl.columns[0].width = Cm(4)
    tbl.columns[1].width = Cm(11)
    for i, h in enumerate(["项目", "内容"]):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.color.rgb = RGBColor(255, 255, 255)
                rr.bold = True
                rr.font.size = Pt(10)
    data = [
        ("程序版本", "XfyunTranscriber v1.0"),
        ("测试音频", "托盘贸易_test.mp3（4.0 MB，约6分钟播客对话）"),
        ("API 服务", "讯飞非实时语音转写大模型 (ifasr_llm)"),
        ("说话人设置", "自动（roleNum=0，盲分模式）"),
    ]
    for i, (k, v) in enumerate(data):
        tbl.rows[i+1].cells[0].text = k
        tbl.rows[i+1].cells[1].text = v
        for j in range(2):
            for pp in tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)

    add_styled_paragraph(doc, "1.3 预期结果 vs 实际结果", bold=True, size=12)

    tbl2 = doc.add_table(rows=2, cols=3, style='Table Grid')
    tbl2.columns[0].width = Cm(2)
    tbl2.columns[1].width = Cm(6)
    tbl2.columns[2].width = Cm(7)
    for i, h in enumerate(["", "修复前（Bug）", "修复后（正确）"]):
        cell = tbl2.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.color.rgb = RGBColor(255, 255, 255)
                rr.bold = True
                rr.font.size = Pt(10)
    tbl2.rows[1].cells[0].text = "说话人数"
    tbl2.rows[1].cells[1].text = "1 人（全部标为「说话人1」）"
    tbl2.rows[1].cells[2].text = "2 人（说话人1 ↔ 说话人2 交替）"
    for j in range(3):
        for pp in tbl2.rows[1].cells[j].paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(10)

    doc.add_paragraph()

    # ════════════════════════════════
    # 二、排查过程
    # ════════════════════════════════
    add_heading_styled(doc, "二、排查过程", level=1)

    add_styled_paragraph(doc, "2.1 第一阶段：追踪说话人分离参数的数据流", bold=True, size=12)
    p = doc.add_paragraph()
    p.add_run("从 GUI 层到 API 层，完整追踪了说话人分离参数的传递路径：").font.size = Pt(10.5)

    steps = [
        "GUI（gui.py）→ 读取 enable_diarization 和 speaker_number → 计算 role_type / role_num",
        "API 客户端（api_client.py）→ 将 roleType/roleNum 加入请求签名参数",
        "轮询管理（poll_manager.py）→ 等待转写完成后返回原始响应",
        "结果解析（result_formatter.py）→ 从 lattice 数据中提取 speaker 字段",
    ]
    for s in steps:
        add_bullet(doc, s)

    add_styled_paragraph(doc, "2.2 第二阶段：发现参数遗漏 Bug", bold=True, size=12)
    p = doc.add_paragraph()
    p.add_run("检查 api_client.py 时发现：当 speaker_number 为「自动」时，role_num 被设为 0，但在签名参数构建时有条件判断 role_num > 0 才添加 roleNum 参数，导致盲分模式下 API 未收到说话人分离指令。").font.size = Pt(10.5)

    add_styled_paragraph(doc, "Bug 代码位置：api_client.py 第 158-160 行（修复前）", bold=True, size=10)
    add_code(doc, "# 修复前：role_num=0 时不发送 roleNum")
    add_code(doc, "if role_num > 0:")
    add_code(doc, "    params[\"roleNum\"] = str(role_num)")

    add_styled_paragraph(doc, "修复：", bold=True, size=10)
    add_code(doc, "# 修复后：roleNum 始终显式发送（0=盲分）")
    add_code(doc, "if role_type > 0:")
    add_code(doc, "    params[\"roleType\"] = str(role_type)")
    add_code(doc, "    params[\"roleNum\"] = str(role_num)")

    add_styled_paragraph(doc, "2.3 第三阶段：编写无头测试脚本验证", bold=True, size=12)
    p = doc.add_paragraph()
    p.add_run("创建独立测试脚本，绕过 GUI 直接调用 API 进行完整转写流程测试。修复 roleNum 发送后，测试结果如下：").font.size = Pt(10.5)

    tbl3 = doc.add_table(rows=3, cols=2, style='Table Grid')
    tbl3.columns[0].width = Cm(4)
    tbl3.columns[1].width = Cm(11)
    for i, h in enumerate(["项目", "结果"]):
        cell = tbl3.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.color.rgb = RGBColor(255, 255, 255)
                rr.bold = True
                rr.font.size = Pt(10)
    tbl3.rows[1].cells[0].text = "上传参数"
    tbl3.rows[1].cells[1].text = "roleType=1, roleNum=0 ✓（正确发送）"
    tbl3.rows[2].cells[0].text = "转写结果"
    tbl3.rows[2].cells[1].text = "56 个片段，但所有 speaker 字段均为 None → 仍只有 1 位说话人"
    for i in range(1, 3):
        for j in range(2):
            for pp in tbl3.rows[i].cells[j].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)

    add_note(doc, "修复了参数发送问题后，说话人分离仍未生效。说明问题还有另一层。", "关键发现")

    add_styled_paragraph(doc, "2.4 第四阶段：诊断 API 返回数据的原始结构", bold=True, size=12)
    p = doc.add_paragraph()
    p.add_run("编写诊断脚本，直接打印 lattice 数组的原始字段结构，定位 speaker 信息在 API 响应中的真实位置。").font.size = Pt(10.5)

    add_styled_paragraph(doc, "诊断结果（前 3 条 lattice 项）：", bold=True, size=10)
    add_code(doc, "Lattice item 0: st keys = ['sc','pa','rt','bg','rl','ed']")
    add_code(doc, "  >>> st.rl = 1")
    add_code(doc, "Lattice item 1: st keys = ['sc','pa','rt','bg','rl','ed']")
    add_code(doc, "  >>> st.rl = 1")
    add_code(doc, "Lattice item 2: st keys = ['sc','pa','rt','bg','rl','ed']")
    add_code(doc, "  >>> st.rl = 2    ← 说话人切换！")

    add_note(doc, "说话人标签存在，但位于 st.rl（role label），而非代码中查找的 json_1best.speaker。", "关键发现")

    add_styled_paragraph(doc, "2.5 第五阶段：定位解析 Bug", bold=True, size=12)
    p = doc.add_paragraph()
    p.add_run("在 result_formatter.py 的 _parse_lattice_format 函数中，大模型 API 的 st 字段为 dict 类型，代码已正确处理了此格式的时间提取（从 st.bg / st.ed 获取毫秒值），但 speaker 字段仅从 json_1best 顶层查找，完全忽略了 st 字典内的 rl 字段。").font.size = Pt(10.5)

    add_styled_paragraph(doc, "Bug 代码位置：result_formatter.py 第 115-119 行（修复前）", bold=True, size=10)
    add_code(doc, "# 修复前：仅查找 json_1best 顶层的 speaker/spk")
    add_code(doc, "speaker = (json_1best.get(\"speaker\", None)")
    add_code(doc, "           or json_1best.get(\"spk\", None))")
    add_code(doc, "# ...大模型API的rl字段被完全忽略...")

    add_styled_paragraph(doc, "修复：", bold=True, size=10)
    add_code(doc, "# 修复后：在 st dict 中提取 rl（role label）")
    add_code(doc, "if isinstance(st_raw, dict):")
    add_code(doc, "    # ...时间提取...")
    add_code(doc, "    if speaker is None:")
    add_code(doc, "        rl = st_raw.get(\"rl\", None)")
    add_code(doc, "        if rl is not None:")
    add_code(doc, "            speaker = str(rl)")

    doc.add_paragraph()

    # ════════════════════════════════
    # 三、根本原因分析
    # ════════════════════════════════
    add_heading_styled(doc, "三、根本原因分析", level=1)

    p = doc.add_paragraph()
    p.add_run("经过完整排查，说话人分离不生效由").font.size = Pt(10.5)
    r = p.add_run("两个独立且相互叠加的 Bug")
    r.bold = True
    r.font.size = Pt(10.5)
    p.add_run("共同导致：").font.size = Pt(10.5)

    add_styled_paragraph(doc, "Bug #1（发送层）：roleNum=0 未发送", bold=True, size=12)

    tbl4 = doc.add_table(rows=5, cols=2, style='Table Grid')
    tbl4.columns[0].width = Cm(3)
    tbl4.columns[1].width = Cm(12)
    rows = [
        ("严重程度", "高 — 直接导致 API 未收到说话人分离指令"),
        ("触发条件", "用户选择「说话人数量 = 自动」时，role_num 被设为 0"),
        ("根本原因", "条件判断 if role_num > 0 排除了 role_num=0 的情况"),
        ("影响范围", "所有选择「自动」说话人数的用户"),
    ]
    for i, (k, v) in enumerate(rows):
        tbl4.rows[i].cells[0].text = k
        tbl4.rows[i].cells[1].text = v
        for j in range(2):
            for pp in tbl4.rows[i].cells[j].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)
    tbl4.rows[4].cells[0].text = "修复方案"
    tbl4.rows[4].cells[1].text = "将 roleNum 的发送逻辑并入 roleType 判断块内，roleType>0 时始终发送 roleNum（0=盲分）"
    for j in range(2):
        for pp in tbl4.rows[4].cells[j].paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(10)

    doc.add_paragraph()

    add_styled_paragraph(doc, "Bug #2（解析层）：rl 字段被忽略，speaker 恒为 None", bold=True, size=12)

    tbl5 = doc.add_table(rows=5, cols=2, style='Table Grid')
    tbl5.columns[0].width = Cm(3)
    tbl5.columns[1].width = Cm(12)
    rows = [
        ("严重程度", "高 — API 返回了说话人信息但代码无法读取"),
        ("触发条件", "所有使用大模型 API（ifasr_llm）的请求"),
        ("根本原因", "大模型 API 的说话人信息存储在 st.rl 字段，但解析代码仅查找 json_1best 顶层的 speaker / spk 字段"),
        ("影响范围", "已修复 Bug #1 后仍无法获得说话人分离结果"),
    ]
    for i, (k, v) in enumerate(rows):
        tbl5.rows[i].cells[0].text = k
        tbl5.rows[i].cells[1].text = v
        for j in range(2):
            for pp in tbl5.rows[i].cells[j].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)
    tbl5.rows[4].cells[0].text = "修复方案"
    tbl5.rows[4].cells[1].text = "在大模型 API 格式（st 为 dict）解析分支中，增加从 st.rl 提取 speaker 的逻辑，优先级低于 json_1best 顶层的 speaker 字段以保持向后兼容"
    for j in range(2):
        for pp in tbl5.rows[4].cells[j].paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(10)

    doc.add_paragraph()

    # ════════════════════════════════
    # 四、数据流分析
    # ════════════════════════════════
    add_heading_styled(doc, "四、数据流分析", level=1)

    add_styled_paragraph(doc, "4.1 请求参数流（修复后）", bold=True, size=12)

    p = doc.add_paragraph()
    p.add_run("用户选择「自动」→ role_type=1, role_num=0 → API 请求 URL 含 ").font.size = Pt(10.5)
    r = p.add_run("roleType=1&roleNum=0")
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    p.add_run(" ✓").font.size = Pt(10.5)

    add_styled_paragraph(doc, "4.2 响应解析流（修复后）", bold=True, size=12)

    steps2 = [
        "API 返回 content.orderResult.lattice[] 数组",
        "每个 lattice item 的 json_1best.st 为 dict，内含 bg / ed / rl 字段",
        "bg=起始帧(ms), ed=结束帧(ms), rl=说话人角色标签(1或2)",
        "result_formatter._parse_lattice_format 识别 st 为 dict",
        "从 st.bg / st.ed 提取时间 → 毫秒转秒",
        "从 st.rl 提取说话人标签 → 覆盖默认的 None / '0'",
        "merge_speaker_segments 合并同一说话人的连续片段",
        "format_output 格式化为「HH:MM:SS - HH:MM:SS 说话人N」",
    ]
    for s in steps2:
        add_bullet(doc, s)

    doc.add_paragraph()

    # ════════════════════════════════
    # 五、API 字段映射对照
    # ════════════════════════════════
    add_heading_styled(doc, "五、API 字段映射对照", level=1)

    add_styled_paragraph(doc, "讯飞两大 API 版本的说话人字段对比：", bold=True, size=12)

    tbl6 = doc.add_table(rows=4, cols=3, style='Table Grid')
    tbl6.columns[0].width = Cm(2.5)
    tbl6.columns[1].width = Cm(6)
    tbl6.columns[2].width = Cm(6.5)
    for i, h in enumerate(["项目", "旧 raasr API", "大模型 API (ifasr_llm)"]):
        cell = tbl6.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.color.rgb = RGBColor(255, 255, 255)
                rr.bold = True
                rr.font.size = Pt(10)
    field_data = [
        ("st 格式", "数值（毫秒）", "dict {bg, ed, rl, rt, sc, pa}"),
        ("说话人字段", "json_1best.speaker", "st.rl（role label）"),
        ("参数名", "pd  /  speaker_number", "roleType  /  roleNum"),
    ]
    for i, (k, v1, v2) in enumerate(field_data):
        tbl6.rows[i+1].cells[0].text = k
        tbl6.rows[i+1].cells[1].text = v1
        tbl6.rows[i+1].cells[2].text = v2
        for j in range(3):
            for pp in tbl6.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)

    add_note(doc, "项目已从旧 raasr API 迁移至大模型 API，但 result_formatter 的解析逻辑未完全适配新格式的 speaker 字段位置。", "注意")

    doc.add_paragraph()

    # ════════════════════════════════
    # 六、修复验证
    # ════════════════════════════════
    add_heading_styled(doc, "六、修复验证", level=1)

    add_styled_paragraph(doc, "6.1 验证方法", bold=True, size=12)
    add_bullet(doc, "编写 Python 测试脚本，完整执行：音频处理 → API 上传 → 轮询 → 结果解析")
    add_bullet(doc, "通过已有的 orderId 重复获取结果，无需重复上传音频")
    add_bullet(doc, "修复后立即重新解析，验证 speaker 数量是否从 1 变为 2")

    add_styled_paragraph(doc, "6.2 验证结果", bold=True, size=12)

    tbl7 = doc.add_table(rows=3, cols=3, style='Table Grid')
    tbl7.columns[0].width = Cm(3)
    tbl7.columns[1].width = Cm(6)
    tbl7.columns[2].width = Cm(6)
    for i, h in enumerate(["指标", "修复前", "修复后"]):
        cell = tbl7.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.color.rgb = RGBColor(255, 255, 255)
                rr.bold = True
                rr.font.size = Pt(10)
    tbl7.rows[1].cells[0].text = "片段数"
    tbl7.rows[1].cells[1].text = "56"
    tbl7.rows[1].cells[2].text = "56"
    tbl7.rows[2].cells[0].text = "说话人数"
    tbl7.rows[2].cells[1].text = "1 人（全部 speaker='0'）"
    tbl7.rows[2].cells[2].text = "2 人（speaker='1' 和 speaker='2'）"
    for i in range(1, 3):
        for j in range(3):
            for pp in tbl7.rows[i].cells[j].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)

    doc.add_paragraph()

    add_styled_paragraph(doc, "6.3 修复后输出示例", bold=True, size=12)
    add_code(doc, "00:00:00 - 00:00:16 说话人1")
    add_code(doc, "Hello大家好，欢迎收听今天的播客...")
    add_code(doc, "")
    add_code(doc, "00:00:16 - 00:00:20 说话人2")
    add_code(doc, "听起来很有意思，啊那我们就直接开始吧...")
    add_code(doc, "")
    add_code(doc, "00:00:20 - 00:00:26 说话人1")
    add_code(doc, "唉咱们先来聊一聊就是这个托盘贸易模式...")

    doc.add_paragraph()

    # ════════════════════════════════
    # 七、修改文件清单
    # ════════════════════════════════
    add_heading_styled(doc, "七、修改文件清单", level=1)

    tbl8 = doc.add_table(rows=4, cols=3, style='Table Grid')
    tbl8.columns[0].width = Cm(4.5)
    tbl8.columns[1].width = Cm(3.5)
    tbl8.columns[2].width = Cm(7)
    for i, h in enumerate(["文件", "修改行", "说明"]):
        cell = tbl8.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "005AA0")
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.color.rgb = RGBColor(255, 255, 255)
                rr.bold = True
                rr.font.size = Pt(10)
    mod_data = [
        ("api_client.py", "158-160", "修复 roleNum=0 时也显式发送参数"),
        ("result_formatter.py", "115-132", "新增大模型 API st.rl 字段的 speaker 提取逻辑"),
        ("result_formatter.py", "85-89", "新增 _log_speaker_stats 日志函数"),
    ]
    for i, (f, line, desc) in enumerate(mod_data):
        tbl8.rows[i+1].cells[0].text = f
        tbl8.rows[i+1].cells[1].text = f"第 {line} 行"
        tbl8.rows[i+1].cells[2].text = desc
        for j in range(3):
            for pp in tbl8.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs:
                    rr.font.size = Pt(10)

    doc.add_paragraph()

    # ════════════════════════════════
    # 八、经验教训
    # ════════════════════════════════
    add_heading_styled(doc, "八、经验教训", level=1)

    lessons = [
        ("API 迁移后需全面审查解析逻辑",
         "从 raasr API 迁移到大模型 API 时，响应结构的 speaker 字段从 json_1best.speaker 迁移到了 st.rl，但解析代码未同步更新。API 迁移时须逐字段比对新旧响应结构差异。"),
        ("参数默认值 ≠ 参数不发送",
         "roleNum=0 表示盲分（自动），与不发送 roleNum 参数不等价。API 的「默认行为」不等同于「参数值=0 的行为」，必须查阅文档确认。"),
        ("分层排查提高效率",
         "采用「发送层 → 响应层 → 解析层」的逐层诊断策略，通过诊断脚本直接查看 API 原始返回数据，快速将问题定位从参数层面缩小到解析层面。"),
        ("日志是排查的关键工具",
         "在解析层添加 speaker 字段的 debug 日志，使得问题在第一次测试中即暴露：所有 speaker=None。后续添加的 _log_speaker_stats 函数可在运行时直观展示说话人统计。"),
    ]
    for i, (title, content) in enumerate(lessons):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(f"教训 {i+1}：{title}")
        r.bold = True
        r.font.size = Pt(11)
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.add_run(content).font.size = Pt(10)

    doc.add_paragraph()

    # ════════════════════════════════
    # 九、结论
    # ════════════════════════════════
    add_heading_styled(doc, "九、结论", level=1)

    p = doc.add_paragraph()
    r = p.add_run("Bug 已排除。")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run("修复涉及两个层面的改动：").font.size = Pt(11)

    add_bullet(doc, "api_client.py — 确保 roleNum 参数始终发送（修复 Bug #1）")
    add_bullet(doc, "result_formatter.py — 从 st.rl 提取大模型 API 的说话人标签（修复 Bug #2）")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("修复后已通过完整终到端测试验证：使用 4MB 播客对话音频（约6分钟）进行转写，成功检测到 2 位说话人，对话段落按说话人交替标注正确。").font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.add_run("后续建议：重新打包 exe（已完成，XfyunTranscriber.exe 85.7 MB）后交付用户，并建议用户在正式使用前用自己的测试音频再次验证说话人分离效果。").font.size = Pt(10.5)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("报告撰写人：AI 编码助手").font.size = Pt(9)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("日期：2025年6月4日").font.size = Pt(9)

    # ── 保存 ──
    output = os.path.join(os.path.dirname(os.path.abspath(__file__)), "说话人分离Bug排除报告.docx")
    doc.save(output)
    print(f"报告已生成: {output}")


if __name__ == "__main__":
    generate()
